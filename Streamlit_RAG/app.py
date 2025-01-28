# Import the libraries
import streamlit as st
import os
import faiss
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_huggingface import HuggingFaceEndpoint
from dotenv import load_dotenv

# st.page config
st.set_page_config(page_title="40-Tech RAG Q&A App", page_icon="🤖")


# Import the Hugging Face API key
load_dotenv()
HUGGING_FACE_API = os.getenv("HUGGING_FACE_API")

# Function to process the input data
def process_input(input_type, input_data):
    """ Process the input data based upon the input type & create the vector store"""
    loader = None
    if input_type == "Web":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            st.error("Invalid PDF file")
            raise ValueError("Invalid PDF file")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            doc = Document(BytesIO(input_data.read()))
        else:
            st.error("Invalid DOCX file")
            raise ValueError("Invalid DOCX file")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    elif input_type == "Text":
        if isinstance(input_data, str):
            documents = input_data
        else:
            st.error("Invalid text input")
            raise ValueError("Invalid text input")
    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode("utf-8")
        elif isinstance(input_data, UploadedFile):
            text = str(input_data.read(), "utf-8")
        else:
            st.error("Invalid TXT file")
            raise ValueError("Invalid TXT file")
        documents = text
    else:
        st.error("Invalid input type")
        raise ValueError("Invalid input type")
    
    # Split the documents into chunks
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == "Web":
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]
    else:
        texts = text_splitter.split_text(documents)
        
    # Create the embeddings
    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}
    
    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )
    # Create FAISS index
    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)
    
    # Create FAISS vector store with the embedding function
    vector_store = FAISS(
        embedding_function=hf_embeddings,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )
    vector_store.add_texts(texts)  # Add documents to the vector store
    return vector_store
    

def answer_question(vectorstore, query):
    """Answers a question based on the provided vectorstore."""
    llm = HuggingFaceEndpoint(repo_id= "meta-llama/Meta-Llama-3-8B-Instruct", 
                              api_key = HUGGING_FACE_API, 
                              temperature= 0.5
                              )
    
    qa = RetrievalQA.from_chain_type(llm=llm, retriever=vectorstore.as_retriever())

    answer = qa({"query": query})
    return answer



def main():
    st.title("40-Tech RAG Q&A App 🚀")
    
    input_type = st.selectbox("Select a source", ["Web", "PDF", "DOCX", "Text", "TXT"])
    if input_type == "Web":
        number_input = st.number_input(min_value=1, max_value=10, step=1, label="Enter the number of URLs")
        input_data = []
        for i in range(number_input):
            url = st.sidebar.text_input(f"Enter the URL of Page:  {i+1}")
            input_data.append(url)
    elif input_type == "PDF":
        input_data = st.file_uploader("Upload a PDF file", type="pdf")
    elif input_type == "DOCX":
        input_data = st.file_uploader("Upload a DOCX file", type=["docx", "doc"])
    elif input_type == "Text":
        input_data = st.text_area("Enter the text")
    elif input_type == "TXT":
        input_data = st.file_uploader("Upload a TXT file", type="txt")
    else:
        st.error("Please select a valid source")
        return

    if st.button("Process"):
        # st.write(process_input(input_type, input_data))
        vectorstore = process_input(input_type, input_data)
        st.session_state["vectorstore"] = vectorstore
    if "vectorstore" in st.session_state:
        query = st.text_input("Ask your question")
        if st.button("Submit"):
            answer = answer_question(st.session_state["vectorstore"], query)
            st.write(answer)


if __name__ == "__main__":
    main()